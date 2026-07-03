import docx
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def create_company_policy_docx():
    doc = docx.Document()
    
    # Title
    doc.add_heading('Acme Corp Employee Handbook & Policies', 0)
    
    # Introduction Section
    doc.add_heading('1. Welcome to Acme Corp', 1)
    doc.add_paragraph(
        "Welcome to Acme Corp! We are thrilled to have you join our team. "
        "This handbook is designed to acquaint you with the company, its policies, "
        "and the benefits and expectations of your employment."
    )
    
    # Working Hours Section
    doc.add_heading('2. Standard Working Hours', 1)
    doc.add_paragraph(
        "Our standard working hours are from 9:00 AM to 5:00 PM, Monday through Friday. "
        "However, we support flexible working schedules and remote work configurations. "
        "Employees should coordinate core working hours with their immediate supervisors. "
        "A minimum of 40 hours per week is expected for full-time employees."
    )
    
    # Annual Leave Section
    doc.add_heading('3. Annual Leave and Paid Time Off (PTO)', 1)
    doc.add_paragraph(
        "Full-time employees receive 20 days of Paid Time Off (PTO) per calendar year, "
        "accrued monthly at a rate of 1.67 days per month. PTO can be used for vacation, "
        "personal time, or illness. Employees must request PTO at least two weeks in advance "
        "for planned vacations, and get manager approval."
    )
    doc.add_paragraph(
        "Unused PTO up to a maximum of 5 days may be carried over into the next calendar year. "
        "Any additional unused PTO exceeding 5 days will be forfeited on December 31st."
    )
    
    # Code of Conduct Section
    doc.add_heading('4. Professional Code of Conduct', 1)
    doc.add_paragraph(
        "Acme Corp is committed to maintaining a safe, inclusive, and professional workspace. "
        "We have zero tolerance for harassment, discrimination, or bullying of any kind. "
        "Employees must treat clients, partners, and colleagues with respect and integrity."
    )
    
    # Save document
    out_path = Path(__file__).resolve().parent / "company_policy.docx"
    doc.save(out_path)
    print(f"Created company policy docx at {out_path}")


def create_ai_handbook_pdf():
    out_path = Path(__file__).resolve().parent / "ai_handbook.pdf"
    doc = SimpleDocTemplate(str(out_path), pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles.byName['Heading1'],
        fontSize=24,
        leading=28,
        spaceAfter=20
    )
    heading_style = ParagraphStyle(
        'HeadingStyle',
        parent=styles.byName['Heading2'],
        fontSize=16,
        leading=20,
        spaceBefore=15,
        spaceAfter=10
    )
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles.byName['Normal'],
        fontSize=11,
        leading=15,
        spaceAfter=10
    )

    story = []
    
    # Title
    story.append(Paragraph("Generative AI & Hardware Acceleration", title_style))
    story.append(Spacer(1, 12))
    
    # Section 1
    story.append(Paragraph("1. Evolution of Large Language Models", heading_style))
    story.append(Paragraph(
        "Large Language Models (LLMs) are deep learning models trained on vast text corpora. "
        "Built on the Transformer architecture introduced by Vaswani et al. in 2017, these models "
        "utilize self-attention mechanisms to understand contextual relationships in language. "
        "Over the years, models have scaled from millions of parameters to hundreds of billions, "
        "enabling emergent capabilities like reasoning, coding, and roleplay.",
        body_style
    ))
    story.append(Spacer(1, 10))
    
    # Section 2
    story.append(Paragraph("2. Llama 3.1: The Open-Weights State of the Art", heading_style))
    story.append(Paragraph(
        "Llama 3.1 is Meta's open-weights model family, offering sizes of 8B, 70B, and 405B parameters. "
        "It supports a context window of 128k tokens, allowing it to digest long documents, textbooks, "
        "or codebases. Llama 3.1 was trained on a dataset of over 15 trillion tokens, focusing on multilingual "
        "capabilities, mathematical reasoning, and tool use. The 8B model is highly popular for local "
        "and low-latency applications due to its high efficiency and strong performance.",
        body_style
    ))
    story.append(Spacer(1, 10))
    
    # Section 3
    story.append(Paragraph("3. Groq LPU: Overcoming the Memory Wall", heading_style))
    story.append(Paragraph(
        "Standard LLM generation is memory-bandwidth bound rather than compute bound, meaning performance "
        "is throttled by how fast weights can be read from memory to the processor. "
        "Groq solved this by designing the Language Processing Unit (LPU), a specialized ASIC designed "
        "specifically for sequential text generation. Unlike GPUs, which rely on parallel processing and high-bandwidth "
        "memory, the Groq LPU utilizes Static Random-Access Memory (SRAM) laid out directly on the chip. "
        "This architectural decision enables deterministic execution and throughputs exceeding 500 tokens per second "
        "for Llama 3.1 8B, making it ideal for real-time interactive agents.",
        body_style
    ))
    
    doc.build(story)
    print(f"Created AI handbook PDF at {out_path}")


if __name__ == "__main__":
    create_company_policy_docx()
    create_ai_handbook_pdf()
